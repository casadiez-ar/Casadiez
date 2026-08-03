from flask import Flask, request, jsonify
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload
import io
import re
import os

app = Flask(__name__)

SCOPES = ["https://www.googleapis.com/auth/drive.file"]

# Colores Casa Diez
COLOR_ANTRACITA = RGBColor(0x2B, 0x2B, 0x2B)  # #2B2B2B
COLOR_DORADO = RGBColor(0xC9, 0xA8, 0x4C)      # #C9A84C


def get_drive_service():
    creds = Credentials(
        None,
        refresh_token=os.environ["GOOGLE_REFRESH_TOKEN"],
        client_id=os.environ["GOOGLE_CLIENT_ID"],
        client_secret=os.environ["GOOGLE_CLIENT_SECRET"],
        token_uri="https://oauth2.googleapis.com/token",
        scopes=SCOPES,
    )
    creds.refresh(Request())
    return build("drive", "v3", credentials=creds)


def set_paragraph_spacing(paragraph, space_before=0, space_after=6, line_spacing=1.15):
    """Aplica interlineado y espaciado a un párrafo."""
    from docx.shared import Pt
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(int(space_before * 20)))
    spacing.set(qn('w:after'), str(int(space_after * 20)))
    spacing.set(qn('w:line'), str(int(line_spacing * 240)))
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)


def extraer_bloques_y_trimestres(contenido):
    """Extrae bloques, contenidos e indicadores del texto generado."""
    import re
    bloques = []
    bloque_actual = None
    contenidos_actuales = []
    indicadores_actuales = []
    en_indicadores = False

    lineas = contenido.split('\n')
    en_criterios = False
    for linea in lineas:
        linea_s = linea.strip()
        # Cuando llega a criterios, guardar el último bloque y detener
        if '## Criterios' in linea_s or 'Criterios de evaluación' in linea_s:
            if bloque_actual:
                bloques.append({
                    'nombre': bloque_actual,
                    'contenidos': contenidos_actuales[:4],
                    'indicadores': indicadores_actuales[:3],
                    'trimestre': ''
                })
                bloque_actual = None
            en_criterios = True
        if en_criterios:
            continue
        if linea_s.startswith('### Bloque') or linea_s.startswith('### bloque'):
            if bloque_actual:
                bloques.append({
                    'nombre': bloque_actual,
                    'contenidos': contenidos_actuales[:4],
                    'indicadores': indicadores_actuales[:3],
                    'trimestre': ''
                })
            bloque_actual = re.sub(r'^###\s*', '', linea_s).strip()
            contenidos_actuales = []
            indicadores_actuales = []
            en_indicadores = False
            en_contenidos = False
        elif '#### Contenidos a trabajar' in linea_s:
            en_contenidos = True
            en_indicadores = False
        elif '#### Indicadores' in linea_s or 'Indicadores de avance' in linea_s:
            en_indicadores = True
            en_contenidos = False
        elif '#### ' in linea_s:
            en_indicadores = False
            en_contenidos = False
        elif linea_s.startswith('- ') and bloque_actual:
            texto = re.sub(r'\*\*', '', linea_s[2:]).strip()[:90]
            if texto and len(texto) > 5:
                if en_indicadores:
                    indicadores_actuales.append(texto)
                elif en_contenidos:
                    contenidos_actuales.append(texto)

    # Guardar último bloque si no llegó a criterios
    if bloque_actual:
        bloques.append({
            'nombre': bloque_actual,
            'contenidos': contenidos_actuales[:4],
            'indicadores': indicadores_actuales[:3],
            'trimestre': ''
        })

    # Asignar trimestres
    en_dist = False
    trimestre_actual = ''
    for linea in lineas:
        linea_s = linea.strip()
        if 'Distribución anual' in linea_s or 'distribucion anual' in linea_s.lower():
            en_dist = True
        if en_dist:
            if '1°' in linea_s and 'trimestre' in linea_s.lower():
                trimestre_actual = '1°'
            elif '2°' in linea_s and 'trimestre' in linea_s.lower():
                trimestre_actual = '2°'
            elif '3°' in linea_s and 'trimestre' in linea_s.lower():
                trimestre_actual = '3°'
            for bloque in bloques:
                nombre_corto = bloque['nombre'].split(':')[0].strip().lower()
                if nombre_corto in linea_s.lower() and trimestre_actual:
                    if not bloque['trimestre']:
                        bloque['trimestre'] = trimestre_actual

    return bloques


def estilo_celda(celda, texto, bold=False, size=9, center=False, fondo=None):
    """Aplica estilo a una celda de tabla."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    p = celda.paragraphs[0]
    p.alignment = 1 if center else 0  # CENTER o LEFT
    run = p.add_run(texto)
    run.font.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = COLOR_ANTRACITA
    run.font.name = 'Arial'
    if fondo:
        tc = celda._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fondo)
        tcPr.append(shd)


def configurar_seccion_landscape(doc):
    """Agrega sección apaisada."""
    sec = doc.add_section()
    sec.orientation = 1
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21.0)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    return sec


def configurar_seccion_portrait(doc):
    """Vuelve a orientación vertical."""
    sec = doc.add_section()
    sec.orientation = 0
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.5)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    return sec


def add_cuadro_simple(doc, bloques):
    """Formato A: Bloque | Contenidos principales | Trimestre."""
    from docx.shared import Cm
    configurar_seccion_landscape(doc)

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = 1
    run_t = p_titulo.add_run('CUADRO DE DISTRIBUCIÓN DE CONTENIDOS')
    run_t.font.bold = True
    run_t.font.size = Pt(12)
    run_t.font.color.rgb = COLOR_ANTRACITA
    run_t.font.name = 'Georgia'

    tabla = doc.add_table(rows=1 + len(bloques), cols=3)
    tabla.style = 'Table Grid'
    tabla.columns[0].width = Cm(7)
    tabla.columns[1].width = Cm(16)
    tabla.columns[2].width = Cm(3)

    for i, enc in enumerate(['BLOQUE', 'CONTENIDOS PRINCIPALES', 'TRIMESTRE']):
        estilo_celda(tabla.rows[0].cells[i], enc, bold=True, size=10, center=True, fondo='E8E8E8')

    for idx, bloque in enumerate(bloques):
        fila = tabla.rows[idx + 1]
        estilo_celda(fila.cells[0], bloque['nombre'], bold=True, size=9)
        contenidos_txt = '\n'.join([f'• {c}' for c in bloque['contenidos']])
        estilo_celda(fila.cells[1], contenidos_txt, size=9)
        estilo_celda(fila.cells[2], bloque.get('trimestre', ''), bold=True, size=10, center=True)

    configurar_seccion_portrait(doc)


def add_cuadro_dc_pba(doc, bloques):
    """Formato B: Bloque | Contenidos | Situaciones/Modos | Indicadores | Tiempo."""
    from docx.shared import Cm
    configurar_seccion_landscape(doc)

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = 1
    run_t = p_titulo.add_run('CUADRO DE DISTRIBUCIÓN DE CONTENIDOS — DC PBA 2018')
    run_t.font.bold = True
    run_t.font.size = Pt(12)
    run_t.font.color.rgb = COLOR_ANTRACITA
    run_t.font.name = 'Georgia'

    tabla = doc.add_table(rows=1 + len(bloques), cols=5)
    tabla.style = 'Table Grid'
    tabla.columns[0].width = Cm(5)
    tabla.columns[1].width = Cm(8)
    tabla.columns[2].width = Cm(7)
    tabla.columns[3].width = Cm(5)
    tabla.columns[4].width = Cm(1.5)

    encabezados = ['BLOQUE', 'CONTENIDOS', 'SITUACIONES DE ENSEÑANZA', 'INDICADORES DE AVANCE', 'TIEMPO']
    for i, enc in enumerate(encabezados):
        estilo_celda(tabla.rows[0].cells[i], enc, bold=True, size=9, center=True, fondo='E8E8E8')

    for idx, bloque in enumerate(bloques):
        fila = tabla.rows[idx + 1]
        estilo_celda(fila.cells[0], bloque['nombre'], bold=True, size=9)
        estilo_celda(fila.cells[1], '\n'.join([f'• {c}' for c in bloque['contenidos']]), size=8)
        estilo_celda(fila.cells[2], 'Ver situaciones de enseñanza en el desarrollo curricular', size=8)
        estilo_celda(fila.cells[3], '\n'.join([f'• {i}' for i in bloque['indicadores']]), size=8)
        estilo_celda(fila.cells[4], bloque.get('trimestre', ''), bold=True, size=9, center=True)

    configurar_seccion_portrait(doc)


def add_cuadro_por_trimestre(doc, bloques):
    """Formato C: Por trimestre con bloques y contenidos dentro de cada uno."""
    from docx.shared import Cm
    configurar_seccion_landscape(doc)

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = 1
    run_t = p_titulo.add_run('DISTRIBUCIÓN DE CONTENIDOS POR TRIMESTRE')
    run_t.font.bold = True
    run_t.font.size = Pt(12)
    run_t.font.color.rgb = COLOR_ANTRACITA
    run_t.font.name = 'Georgia'

    # Agrupar bloques por trimestre
    por_trimestre = {'1°': [], '2°': [], '3°': []}
    for bloque in bloques:
        t = bloque.get('trimestre', '')
        if t in por_trimestre:
            por_trimestre[t].append(bloque)

    tabla = doc.add_table(rows=1, cols=3)
    tabla.style = 'Table Grid'
    tabla.columns[0].width = Cm(8.5)
    tabla.columns[1].width = Cm(8.5)
    tabla.columns[2].width = Cm(8.5)

    for i, trim in enumerate(['1° TRIMESTRE', '2° TRIMESTRE', '3° TRIMESTRE']):
        estilo_celda(tabla.rows[0].cells[i], trim, bold=True, size=10, center=True, fondo='E8E8E8')

    # Agregar fila de contenido
    fila = tabla.add_row()
    for i, key in enumerate(['1°', '2°', '3°']):
        celda = fila.cells[i]
        bloques_trim = por_trimestre.get(key, [])
        if bloques_trim:
            texto = ''
            for b in bloques_trim:
                texto += f"{b['nombre']}\n"
                for c in b['contenidos']:
                    texto += f"  • {c}\n"
                texto += '\n'
            p = celda.paragraphs[0]
            run = p.add_run(texto.strip())
            run.font.size = Pt(8)
            run.font.color.rgb = COLOR_ANTRACITA
            run.font.name = 'Arial'
        else:
            estilo_celda(celda, 'A definir', size=8, center=True)

    configurar_seccion_portrait(doc)


def add_cuadro_contenidos(doc, contenido, formato='simple'):
    """Genera el cuadro de distribución de contenidos según el formato elegido."""
    bloques = extraer_bloques_y_trimestres(contenido)
    if not bloques:
        return
    if formato == 'dc_pba':
        add_cuadro_dc_pba(doc, bloques)
    elif formato == 'por_trimestre':
        add_cuadro_por_trimestre(doc, bloques)
    else:
        add_cuadro_simple(doc, bloques)


def add_heading_styled(doc, text, level):
    """Agrega un título con color antracita."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.color.rgb = COLOR_ANTRACITA
    run.font.bold = True

    if level == 1:
        run.font.name = 'Georgia'
        run.font.size = Pt(16)
        set_paragraph_spacing(p, space_before=12, space_after=6)
    elif level == 2:
        run.font.name = 'Georgia'
        run.font.size = Pt(14)
        set_paragraph_spacing(p, space_before=10, space_after=4)
    elif level == 3:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.italic = True
        set_paragraph_spacing(p, space_before=8, space_after=3)

    return p


def add_caratula(doc, datos):
    """Agrega la carátula con datos del docente y símbolo Casa Diez."""
    # Salto de sección para que la carátula sea página propia
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)

    # Símbolo 匠心 centrado arriba en dorado
    p_simbolo = doc.add_paragraph()
    p_simbolo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_simbolo = p_simbolo.add_run('匠心')
    run_simbolo.font.size = Pt(36)
    run_simbolo.font.color.rgb = COLOR_DORADO
    run_simbolo.font.name = 'Georgia'
    set_paragraph_spacing(p_simbolo, space_before=60, space_after=4)

    # Nombre Casa Diez debajo del símbolo
    p_nombre = doc.add_paragraph()
    p_nombre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_nombre = p_nombre.add_run('Casa Diez')
    run_nombre.font.size = Pt(13)
    run_nombre.font.color.rgb = COLOR_DORADO
    run_nombre.font.name = 'Arial'
    set_paragraph_spacing(p_nombre, space_before=0, space_after=40)

    # Línea separadora dorada
    p_linea = doc.add_paragraph()
    p_linea.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_linea = p_linea.add_run('────────────────────────')
    run_linea.font.color.rgb = COLOR_DORADO
    run_linea.font.size = Pt(10)
    set_paragraph_spacing(p_linea, space_before=0, space_after=30)

    # Título principal
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = p_titulo.add_run('PLANIFICACIÓN ANUAL')
    run_titulo.font.size = Pt(22)
    run_titulo.font.bold = True
    run_titulo.font.color.rgb = COLOR_ANTRACITA
    run_titulo.font.name = 'Georgia'
    set_paragraph_spacing(p_titulo, space_before=0, space_after=8)

    # Subtítulo — grado y sección (sin materia)
    grado_raw = datos.get('grado', '')
    seccion = datos.get('seccion', '')
    import re as _re2
    grado_limpio = _re2.sub(r'(?i)\bgrado\b', '', grado_raw)
    grado_limpio = _re2.sub(r'(?i)\bsecci[oó]n\b', '', grado_limpio)
    grado_limpio = grado_limpio.replace('"', '').replace("'", '').replace(',', '').replace('-', '').strip()
    grado_limpio = _re2.sub(r'\s+', ' ', grado_limpio).strip()

    # Detectar primaria o secundaria por cargo
    cargo = datos.get('cargo', '')
    if 'grado' in cargo.lower():
        nivel_word = 'grado'
    else:
        nivel_word = 'año'

    if grado_limpio:
        subtitulo_text = f'{grado_limpio} {nivel_word}' if not seccion else f'{grado_limpio} {nivel_word} {seccion}'
    else:
        subtitulo_text = f'Sección {seccion}' if seccion else ''

    # Subtitulo: insertar nivel_word entre numero y seccion
    # grado_limpio puede ser "3 A" -- insertar nivel_word antes de la seccion
    partes = grado_limpio.split()
    if len(partes) >= 2:
        subtitulo_text = partes[0] + " " + nivel_word + " " + " ".join(partes[1:])
    elif len(partes) == 1:
        subtitulo_text = partes[0] + " " + nivel_word
    else:
        subtitulo_text = nivel_word

    if subtitulo_text:
        p_subtitulo = doc.add_paragraph()
        p_subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = p_subtitulo.add_run(subtitulo_text)
        run_sub.font.size = Pt(14)
        run_sub.font.color.rgb = COLOR_ANTRACITA
        run_sub.font.name = 'Georgia'
        run_sub.font.italic = True
        set_paragraph_spacing(p_subtitulo, space_before=0, space_after=50)

    # Segunda línea separadora
    p_linea2 = doc.add_paragraph()
    p_linea2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_linea2 = p_linea2.add_run('────────────────────────')
    run_linea2.font.color.rgb = COLOR_DORADO
    run_linea2.font.size = Pt(10)
    set_paragraph_spacing(p_linea2, space_before=0, space_after=30)

    # Datos del docente centrados
    campos = [
        ('Establecimiento', datos.get('establecimiento', '')),
        ('Docente', datos.get('docente', '')),
        ('Cargo', datos.get('cargo', '')),
        ('Ciclo lectivo', datos.get('ciclo', '')),
    ]
    for campo, valor in campos:
        if valor:
            p_dato = doc.add_paragraph()
            p_dato.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_campo = p_dato.add_run(f'{campo}: ')
            run_campo.font.bold = True
            run_campo.font.size = Pt(11)
            run_campo.font.color.rgb = COLOR_ANTRACITA
            run_campo.font.name = 'Arial'
            run_valor = p_dato.add_run(valor)
            run_valor.font.size = Pt(11)
            run_valor.font.color.rgb = COLOR_ANTRACITA
            run_valor.font.name = 'Arial'
            set_paragraph_spacing(p_dato, space_before=0, space_after=6)

    # Salto de página al terminar la carátula
    doc.add_page_break()


def quitar_bordes_tabla(tabla):
    """Quita todos los bordes visibles de una tabla."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tbl = tabla._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def add_encabezado(doc, establecimiento, docente):
    """Agrega encabezado con tabla 2 columnas: establecimiento izquierda, docente derecha."""
    from docx.shared import Inches
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    section = doc.sections[-1]
    header = section.header
    header.is_linked_to_previous = False

    # Limpiar contenido existente
    for p in header.paragraphs:
        p.clear()

    # Crear tabla de 2 columnas sin bordes — ancho = área de texto exacta
    tabla = header.add_table(rows=1, cols=2, width=Inches(6.1))
    tabla.autofit = False
    tabla.columns[0].width = Inches(3.05)
    tabla.columns[1].width = Inches(3.05)
    quitar_bordes_tabla(tabla)

    # Celda izquierda — establecimiento
    celda_izq = tabla.cell(0, 0)
    p_izq = celda_izq.paragraphs[0]
    p_izq.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_est = p_izq.add_run(establecimiento)
    run_est.font.size = Pt(9)
    run_est.font.color.rgb = COLOR_ANTRACITA
    run_est.font.name = 'Arial'

    # Celda derecha — docente
    celda_der = tabla.cell(0, 1)
    p_der = celda_der.paragraphs[0]
    p_der.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_doc = p_der.add_run(docente)
    run_doc.font.size = Pt(9)
    run_doc.font.color.rgb = COLOR_ANTRACITA
    run_doc.font.name = 'Arial'

    # Línea dorada debajo — en el párrafo vacío que queda después de la tabla
    p_linea = header.add_paragraph()
    pPr = p_linea._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C9A84C')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_pie_pagina(doc, ciclo):
    """Agrega pie de página con tabla 2 columnas: ciclo izquierda, página derecha."""
    from docx.shared import Inches
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    section = doc.sections[-1]
    footer = section.footer
    footer.is_linked_to_previous = False

    for p in footer.paragraphs:
        p.clear()

    # Tabla de 2 columnas sin bordes
    tabla = footer.add_table(rows=1, cols=2, width=Inches(6.1))
    tabla.autofit = False
    tabla.columns[0].width = Inches(3.05)
    tabla.columns[1].width = Inches(3.05)
    quitar_bordes_tabla(tabla)

    # Celda izquierda — ciclo lectivo
    celda_izq = tabla.cell(0, 0)
    p_izq = celda_izq.paragraphs[0]
    p_izq.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_ciclo = p_izq.add_run(f'Ciclo lectivo {ciclo}')
    run_ciclo.font.size = Pt(9)
    run_ciclo.font.color.rgb = COLOR_ANTRACITA
    run_ciclo.font.name = 'Arial'

    # Celda derecha — número de página
    celda_der = tabla.cell(0, 1)
    p_der = celda_der.paragraphs[0]
    p_der.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run_pag = p_der.add_run('Página ')
    run_pag.font.size = Pt(9)
    run_pag.font.color.rgb = COLOR_ANTRACITA
    run_pag.font.name = 'Arial'

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run_num = p_der.add_run()
    run_num._r.append(fldChar1)
    run_num._r.append(instrText)
    run_num._r.append(fldChar2)
    run_num.font.size = Pt(9)
    run_num.font.color.rgb = COLOR_ANTRACITA


def extraer_datos(contenido):
    """Extrae datos de identificación del contenido para la carátula y encabezado."""
    datos = {
        'establecimiento': '',
        'docente': '',
        'cargo': '',
        'ciclo': '',
        'materia': '',
        'grado': '',
        'seccion': '',
        'formato_cuadro': 'simple',
    }
    patrones = {
        'establecimiento': r'\*\*Establecimiento:\*\*\s*(.+)',
        'docente': r'\*\*Docente:\*\*\s*(.+)',
        'cargo': r'\*\*Cargo:\*\*\s*(.+)',
        'ciclo': r'\*\*Ciclo lectivo:\*\*\s*(.+)',
        'materia': r'\*\*Materia:\*\*\s*(.+)',
        'grado': r'\*\*Año/Grado y Sección:\*\*\s*(.+)',
    }
    for campo, patron in patrones.items():
        match = re.search(patron, contenido)
        if match:
            datos[campo] = match.group(1).strip()
    return datos


def limpiar_nombre_archivo(texto):
    """Limpia un texto para usarlo como nombre de archivo."""
    texto = texto.strip()
    texto = re.sub(r'[^\w\s\-]', '', texto, flags=re.UNICODE)
    texto = re.sub(r'\s+', '_', texto)
    texto = re.sub(r'_+', '_', texto)
    return texto


@app.route('/generar', methods=['POST'])
def generar():
    data = request.get_json()
    titulo = data.get('titulo', 'Planificacion_Anual')
    contenido = data.get('contenido', '')
    folder_id = os.environ.get('DRIVE_FOLDER_ID', '')

    # Extraer datos para carátula y encabezado
    datos = extraer_datos(contenido)

    doc = Document()

    # Configurar estilo base
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.font.color.rgb = COLOR_ANTRACITA

    # Márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # Agregar carátula
    add_caratula(doc, datos)

    # Agregar nueva sección para el contenido (así el encabezado/pie
    # solo aparece desde la página 2 en adelante)
    nueva_seccion = doc.add_section()
    nueva_seccion.header_distance = Cm(1.25)
    nueva_seccion.footer_distance = Cm(1.25)
    # Forzar numeración desde 1 con XML directo
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:start'), '1')
    nueva_seccion._sectPr.append(pgNumType)

    # Agregar encabezado y pie de página (solo en sección del contenido)
    add_encabezado(
        doc,
        datos.get('establecimiento', 'Casa Diez'),
        datos.get('docente', '')
    )
    ciclo_pie = datos.get('ciclo', '').strip()
    if not ciclo_pie:
        ciclo_pie = '2026'
    add_pie_pagina(doc, ciclo_pie)

    # Procesar contenido línea por línea
    # El cuadro se insertará automáticamente después de las expectativas de logro
    lineas = contenido.split('\n')

    for linea in lineas:
        linea_strip = linea.strip()

        if not linea_strip:
            p = doc.add_paragraph()
            set_paragraph_spacing(p, space_before=0, space_after=3)
            continue

        # ## Título nivel 2
        if linea_strip.startswith('## '):
            texto = linea_strip[3:]
            # Insertar cuadro de contenidos antes del desarrollo curricular
            if 'Desarrollo curricular' in texto or 'desarrollo curricular' in texto.lower():
                formato_cuadro = datos.get('formato_cuadro', 'simple')
                if formato_cuadro == 'Por bloque con situaciones de enseñanza':
                    add_cuadro_contenidos(doc, contenido, formato='dc_pba')
                elif formato_cuadro == 'Por trimestre con bloques dentro':
                    add_cuadro_contenidos(doc, contenido, formato='por_trimestre')
                else:
                    add_cuadro_contenidos(doc, contenido, formato='simple')
            add_heading_styled(doc, texto, level=2)
            continue

        # ### Título nivel 3
        if linea_strip.startswith('### '):
            texto = linea_strip[4:]
            add_heading_styled(doc, texto, level=3)
            continue

        # #### Subtítulo nivel 4
        if linea_strip.startswith('#### '):
            texto = linea_strip[5:]
            p = doc.add_paragraph()
            run = p.add_run(texto)
            run.bold = True
            run.italic = True
            run.font.size = Pt(11)
            run.font.color.rgb = COLOR_ANTRACITA
            run.font.name = 'Arial'
            set_paragraph_spacing(p, space_before=6, space_after=3)
            continue

        # - Lista con bullet
        if linea_strip.startswith('- '):
            texto = linea_strip[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            partes = re.split(r'\*\*', texto)
            for i, parte in enumerate(partes):
                run = p.add_run(parte)
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.font.color.rgb = COLOR_ANTRACITA
                if i % 2 == 1:
                    run.bold = True
            set_paragraph_spacing(p, space_before=0, space_after=3)
            continue

        # Párrafo con negrita **texto**
        if '**' in linea_strip:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            partes = re.split(r'\*\*', linea_strip)
            for i, parte in enumerate(partes):
                run = p.add_run(parte)
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.font.color.rgb = COLOR_ANTRACITA
                if i % 2 == 1:
                    run.bold = True
            set_paragraph_spacing(p, space_before=0, space_after=4)
            continue

        # Párrafo de texto normal
        p = doc.add_paragraph(linea_strip)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.font.color.rgb = COLOR_ANTRACITA
        set_paragraph_spacing(p, space_before=0, space_after=4)

    # Guardar en buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Nombre del archivo: sin materia, sin "grado", solo grado numerico y sección
    import re as _re
    grado_raw_arch = datos.get('grado', '').strip()
    seccion_arch = datos.get('seccion', '').strip()

    grado_arch = _re.sub(r'(?i)\bgrado\b', '', grado_raw_arch)
    grado_arch = _re.sub(r'(?i)\bsecci[oó]n\b', '', grado_arch)
    grado_arch = grado_arch.replace('"', '').replace("'", '').replace(',', '').replace('-', '').strip()
    grado_arch = _re.sub(r'\s+', ' ', grado_arch).strip()

    if grado_arch and seccion_arch:
        nombre_archivo = f"Planificacion Anual {grado_arch} {seccion_arch}.docx"
    elif grado_arch:
        nombre_archivo = f"Planificacion Anual {grado_arch}.docx"
    else:
        nombre_archivo = "Planificacion Anual.docx"

    # Subir a Google Drive
    service = get_drive_service()

    file_metadata = {
        'name': nombre_archivo,
        'parents': [folder_id]
    }

    media = MediaIoBaseUpload(
        buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

    file = service.files().create(
        body=file_metadata,
        media_body=media,
        fields='id, webViewLink'
    ).execute()

    return jsonify({
        'status': 'ok',
        'documentId': file.get('id'),
        'url': file.get('webViewLink'),
        'nombre_archivo': nombre_archivo
    })


@app.route('/')
def health():
    return 'Casa Diez - Generador de Planificaciones OK'


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8080)
